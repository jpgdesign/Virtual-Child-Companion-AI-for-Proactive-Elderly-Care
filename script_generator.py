# -*- coding: utf-8 -*-
"""
script_generator.py
劇本生成模組 - 專門與奶奶對話版本

功能：
1. 生成12組基礎劇本（3個話題來源 × 4個目的槽）
2. 嚴格限定對話對象為奶奶
3. 保存到資料夾
4. 以日期區分不同訓練批次
"""

import openai
import json
import os
import pandas as pd
from datetime import datetime
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional, Tuple
from docx import Document
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# API配置 - 直接輸入
def setup_openai():
    """設置OpenAI API"""
    openai_key = os.getenv("OPENAI_API_KEY")

    if not openai_key:
        raise Exception("Missing OPENAI_API_KEY environment variable.")

    from openai import OpenAI
    return OpenAI(api_key=openai_key)

client = setup_openai()

@dataclass
class DialogueStep:
    """單步對話 - 專門與奶奶對話"""
    step_number: int
    child_dialogue: str
    expected_grandma_response: str  # 改為 grandma_response

@dataclass
class Script:
    """劇本結構 - 奶奶對話版本"""
    script_id: int
    script_type: str  # "奶奶對話劇本"
    source: str
    target_slot: str
    total_steps: int
    steps: List[DialogueStep]
    target_info: List[str]

class ScriptGenerator:
    """劇本生成器 - 專門與奶奶對話"""

    def __init__(self, data_path="data/"):
        self.data_path = data_path
        self.base_output_path = "outputs/"

        # 配置 - 根據實際文件更新
        self.sources = ["背景資訊", "喜好興趣", "作息"]
        self.target_slots = ["用藥狀況", "睡眠狀態", "作息活動", "飲食狀況"]

        # 目標資訊定義
        self.target_items_map = {
            "用藥狀況": ["量血壓情況", "服藥情況", "身體不適", "用藥時間"],
            "睡眠狀態": ["起床時間", "睡眠時間", "小睡情況", "睡眠品質"],
            "作息活動": ["上廁所情況", "居家運動", "外出情況", "洗澡情況", "喝水情況", "看電視情況"],
            "飲食狀況": ["三餐時間", "食物內容", "廚房使用", "冰箱使用"]
        }

        # 延後創建輸出路徑，避免每次初始化都創建新資料夾
        self.session_folder = None
        self.script_path = None

        # 進度追蹤
        self.progress_file = None
        self.generated_scripts = []

        logger.info("📝 劇本生成器初始化完成 - 奶奶對話專用版")

    def _setup_output_paths(self):
        """設置輸出路徑（只在實際生成時才創建）"""
        if self.session_folder is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            self.session_folder = os.path.join(self.base_output_path, "scripts", f"grandma_session_{timestamp}")
            self.script_path = os.path.join(self.session_folder, "奶奶劇本")

            # 確保輸出目錄存在
            os.makedirs(self.script_path, exist_ok=True)

            # 設置進度追蹤檔案
            self.progress_file = os.path.join(self.session_folder, "progress.json")

            logger.info(f"📁 奶奶劇本路徑: {self.script_path}")
            logger.info(f"📊 進度追蹤檔案: {self.progress_file}")

    def _save_progress(self):
        """保存當前進度"""
        if self.progress_file:
            progress_data = {
                "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S'),
                "dialogue_target": "奶奶",  # 明確標記對話對象
                "total_scripts": len(self.generated_scripts),
                "scripts": [asdict(script) for script in self.generated_scripts],
                "completed_combinations": []
            }

            # 記錄已完成的組合
            for script in self.generated_scripts:
                progress_data["completed_combinations"].append({
                    "source": script.source,
                    "target_slot": script.target_slot,
                    "script_id": script.script_id
                })

            with open(self.progress_file, 'w', encoding='utf-8') as f:
                json.dump(progress_data, f, ensure_ascii=False, indent=2)

            logger.info(f"💾 進度已保存: {len(self.generated_scripts)}/12 (奶奶對話)")

    def _load_progress(self) -> bool:
        """載入之前的進度（如果存在）"""
        if self.progress_file and os.path.exists(self.progress_file):
            try:
                with open(self.progress_file, 'r', encoding='utf-8') as f:
                    progress_data = json.load(f)

                # 重建劇本物件
                for script_data in progress_data.get("scripts", []):
                    steps = [DialogueStep(**step) for step in script_data["steps"]]
                    script_data["steps"] = steps
                    self.generated_scripts.append(Script(**script_data))

                logger.info(f"📂 載入之前進度: {len(self.generated_scripts)}/12 (奶奶對話)")
                return True

            except Exception as e:
                logger.warning(f"載入進度失敗: {e}")
                return False
        return False

    def _get_next_script_id(self) -> int:
        """獲取下一個劇本ID"""
        if not self.generated_scripts:
            return 1
        return max(script.script_id for script in self.generated_scripts) + 1

    def _is_combination_completed(self, source: str, target_slot: str) -> bool:
        """檢查指定組合是否已完成"""
        for script in self.generated_scripts:
            if script.source == source and script.target_slot == target_slot:
                return True
        return False

    def load_source_content(self) -> Dict[str, str]:
        """從data資料夾直接載入話題來源檔案"""
        source_content = {}

        try:
            # 確保輸出目錄存在
            os.makedirs(self.base_output_path, exist_ok=True)
            
            # 直接讀取背景資訊
            background_path = os.path.join(self.data_path, "01_背景資訊.docx")
            if not os.path.exists(background_path):
                raise FileNotFoundError(f"找不到檔案: {background_path}")
            doc = Document(background_path)
            source_content["背景資訊"] = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            logger.info(f"✅ 載入背景資訊: ({len(source_content['背景資訊'])} 字元)")

            # 直接讀取喜好興趣
            interest_path = os.path.join(self.data_path, "03_喜好興趣.docx")
            if not os.path.exists(interest_path):
                raise FileNotFoundError(f"找不到檔案: {interest_path}")
            doc = Document(interest_path)
            source_content["喜好興趣"] = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            logger.info(f"✅ 載入喜好興趣: ({len(source_content['喜好興趣'])} 字元)")

            # 直接讀取作息
            schedule_path = os.path.join(self.data_path, "02_作息.xlsx")
            if not os.path.exists(schedule_path):
                raise FileNotFoundError(f"找不到檔案: {schedule_path}")
            df = pd.read_excel(schedule_path)
            source_content["作息"] = df.to_string(index=False)
            logger.info(f"✅ 載入作息: ({len(source_content['作息'])} 字元)")

            logger.info("✅ 所有話題來源內容載入完成")
            return source_content

        except Exception as e:
            logger.error(f"載入話題來源失敗: {e}")
            raise Exception(f"無法載入必要的話題來源檔案: {e}")

    def generate_all_scripts(self) -> List[Script]:
        """生成所有12組劇本，支援斷點續傳 - 專門與奶奶對話"""
        logger.info("🎭 開始生成12組奶奶對話劇本...")

        # 先設置輸出路徑
        self._setup_output_paths()

        # 嘗試載入之前的進度
        self._load_progress()

        source_content = self.load_source_content()
        script_id = self._get_next_script_id()

        # 生成12組劇本
        logger.info("📝 生成12組奶奶對話劇本...")

        for source in self.sources:
            for target_slot in self.target_slots:
                # 檢查是否已經完成
                if self._is_combination_completed(source, target_slot):
                    logger.info(f"⏭️ 跳過已完成的劇本: {source} → {target_slot}")
                    continue

                try:
                    logger.info(f"🎯 正在生成奶奶對話劇本{script_id}: {source} → {target_slot}")
                    script = self._generate_script(
                        script_id, source, target_slot, source_content[source]
                    )
                    self.generated_scripts.append(script)

                    # 立即保存進度
                    self._save_progress()

                    logger.info(f"✅ 奶奶對話劇本{script_id}生成成功")
                    script_id += 1

                except Exception as e:
                    logger.error(f"❌ 奶奶對話劇本{script_id}生成失敗: {e}")
                    logger.error(f"詳細錯誤: {str(e)}")
                    # 保存當前進度後拋出異常
                    self._save_progress()
                    raise Exception(f"奶奶對話劇本{script_id}生成失敗，進度已保存，請檢查GPT API設置或網路連線: {e}")

        logger.info(f"🎉 成功生成 {len(self.generated_scripts)} 個奶奶對話劇本")
        return self.generated_scripts

    def _generate_script(self, script_id: int, source: str, target_slot: str, source_content: str) -> Script:
        """生成劇本 - 專門與奶奶對話"""
        target_info = self.target_items_map[target_slot]

        prompt = f"""請你扮演**孫女**，以**自然、生活化、情感真摯的語氣**，與一位**奶奶**聊天。

**重要約束：**
- 對話對象必須是**奶奶**（女性長輩）
- 絕對不可以出現爺爺、老爺爺等男性稱呼
- 所有回應都要符合奶奶的身份和語氣
- 使用適合與奶奶對話的溫馨、親密語調

我的目的是**隱性探查健康行為**，但**不可以直接問「今天吃藥了沒」、「有沒有量血壓」這類顯性問題**，必須透過分享、閒聊、帶情感的引導方式進行。每句話不需都是問句，可自然穿插**分享個人生活、關注興趣、誇獎或懷舊話題**來帶出回應。

話題來源：{source}
來源內容：{source_content[:500]}...

目的槽：{target_slot}
需要收集的資訊：{', '.join(target_info)}

要求：
1. 對話步驟數：5-10步
2. 每步包含「孫女對話」和「預期奶奶回答」
3. 對話要自然流暢，從話題來源漸進引導到目的槽
4. 不要直接詢問目標資訊，要使用間接、關心的方式
5. 孫女語氣要溫暖親切，適合與奶奶對話
6. 預期奶奶回答要具體且包含目標資訊，符合奶奶的說話方式
7. **嚴格確保所有回應都是奶奶的身份，不可出現爺爺等男性角色**

**請嚴格按照以下格式輸出：**

第1步：
孫女：[自然生活化的開場，適合與奶奶對話]
預期奶奶：[包含相關資訊的溫暖回應，符合奶奶語氣]

第2步：
孫女：[分享或關心式的引導，適合與奶奶互動]
預期奶奶：[包含目標資訊的具體回應，奶奶的說話風格]

第3步：
孫女：[繼續引導]
預期奶奶：[具體回應]

依此類推至第5-10步...

**注意：請確保每一步都包含「孫女：」和「預期奶奶：」標記，絕對不可出現爺爺相關內容**
        """

        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "你是專業的對話劇本設計師，專門為老人關懷對話設計自然溫馨的劇本。你必須嚴格確保所有對話都是與奶奶（女性長輩）進行，絕對不可以出現爺爺、老爺爺等男性角色。請嚴格按照指定格式輸出劇本。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )

            content = response.choices[0].message.content
            logger.info(f"GPT回應內容預覽: {content[:200]}...")

            # 檢查是否包含不當內容
            if self._contains_unwanted_male_references(content):
                logger.warning("⚠️ 檢測到男性角色，重新生成...")
                # 可以選擇重新生成或拋出異常
                raise Exception("生成的劇本包含不當的男性角色，需要重新生成")

            steps = self._parse_script_content(content)

            # 檢查步驟數量
            if len(steps) < 3:
                raise Exception(f"生成的對話步驟太少: {len(steps)}步")

            return Script(
                script_id=script_id,
                script_type="奶奶對話劇本",
                source=source,
                target_slot=target_slot,
                total_steps=len(steps),
                steps=steps,
                target_info=target_info
            )

        except Exception as e:
            logger.error(f"GPT生成失敗: {e}")
            raise Exception(f"無法生成奶奶對話劇本: {e}")

    def _contains_unwanted_male_references(self, content: str) -> bool:
        """檢查內容是否包含不當的男性角色引用"""
        unwanted_terms = [
            "爺爺", "老爺爺", "阿公", "公公", "外公",
            "老公公", "老先生", "大爺", "老爺",
            "爺爺：", "老爺爺：", "阿公：", "公公：", "外公："
        ]

        content_lower = content.lower()
        for term in unwanted_terms:
            if term in content or term.lower() in content_lower:
                logger.warning(f"發現不當男性引用: {term}")
                return True
        return False

    def _parse_script_content(self, content: str) -> List[DialogueStep]:
        """解析劇本內容 - 增強版本，支援多種格式，專門處理奶奶對話"""
        steps = []
        lines = content.split('\n')
        current_step = None
        child_dialogue = ""
        expected_response = ""

        for line in lines:
            line = line.strip()

            # 尋找步驟標記（支援多種格式）
            step_patterns = [
                r'第(\d+)步',
                r'步驟(\d+)',
                r'(\d+)\.',
                r'Step\s*(\d+)',
                r'第(\d+)輪'
            ]

            step_found = False
            for pattern in step_patterns:
                import re
                match = re.search(pattern, line)
                if match:
                    # 保存前一步
                    if current_step is not None and child_dialogue and expected_response:
                        steps.append(DialogueStep(current_step, child_dialogue, expected_response))

                    # 開始新的步驟
                    current_step = int(match.group(1))
                    child_dialogue = ""
                    expected_response = ""
                    step_found = True
                    break

            if step_found:
                continue

            # 尋找孫女對話（支援多種格式）
            child_patterns = [
                '孫女：', '孫女:', '孫女 :', '孫女 ：',
                '小孩：', '小孩:', '兒童：', '兒童:',
                '女孩：', '女孩:', '女：', '女:'
            ]

            for pattern in child_patterns:
                if line.startswith(pattern):
                    child_dialogue = line.split(pattern, 1)[1].strip()
                    break

            # 尋找奶奶回應（專門針對奶奶的格式）
            grandma_patterns = [
                '預期奶奶：', '預期奶奶:', '奶奶：', '奶奶:',
                '阿嬤：', '阿嬤:', '外婆：', '外婆:',
                '預期回應：', '預期回應:', '回應：', '回應:',
                '長者：', '長者:'  # 保留中性稱呼
            ]

            for pattern in grandma_patterns:
                if line.startswith(pattern):
                    expected_response = line.split(pattern, 1)[1].strip()
                    break

        # 保存最後一步
        if current_step is not None and child_dialogue and expected_response:
            steps.append(DialogueStep(current_step, child_dialogue, expected_response))

        # 如果解析失敗，記錄詳細信息
        if not steps:
            logger.error("奶奶對話劇本解析失敗的內容:")
            logger.error("=" * 50)
            logger.error(content)
            logger.error("=" * 50)
            raise Exception("奶奶對話劇本解析失敗，無法提取有效對話步驟。GPT回應格式可能不正確。")

        logger.info(f"✅ 成功解析 {len(steps)} 個奶奶對話步驟")
        return steps

    def save_scripts(self, scripts: List[Script]) -> str:
        """保存劇本 - 奶奶對話版本"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        # 保存JSON格式
        scripts_data = {
            "generation_time": timestamp,
            "script_type": "奶奶對話劇本",
            "dialogue_target": "奶奶",
            "total_scripts": len(scripts),
            "scripts": [asdict(script) for script in scripts]
        }

        json_filename = os.path.join(self.script_path, f"奶奶對話劇本_{timestamp}.json")
        with open(json_filename, 'w', encoding='utf-8') as f:
            json.dump(scripts_data, f, ensure_ascii=False, indent=2)

        # 保存可讀格式
        txt_filename = os.path.join(self.script_path, f"奶奶對話劇本_readable_{timestamp}.txt")
        with open(txt_filename, 'w', encoding='utf-8') as f:
            f.write(f"🎭 奶奶對話劇本\n")
            f.write("=" * 80 + "\n")
            f.write(f"生成時間: {timestamp}\n")
            f.write(f"對話對象: 奶奶\n")
            f.write(f"劇本總數: {len(scripts)}\n\n")

            for script in scripts:
                self._write_script_to_file(f, script)

        # 保存整體統計資訊
        summary_file = os.path.join(self.session_folder, f"grandma_session_summary_{timestamp}.json")
        summary_data = {
            "session_timestamp": timestamp,
            "dialogue_target": "奶奶",
            "session_folder": self.session_folder,
            "scripts": {
                "count": len(scripts),
                "file_path": json_filename
            },
            "total_scripts": len(scripts),
            "sources": list(set(s.source for s in scripts)),
            "target_slots": list(set(s.target_slot for s in scripts))
        }

        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary_data, f, ensure_ascii=False, indent=2)

        logger.info(f"💾 奶奶對話劇本保存完成:")
        logger.info(f"  劇本檔案: {json_filename}")
        logger.info(f"  可讀檔案: {txt_filename}")
        logger.info(f"  會話總結: {summary_file}")

        return json_filename

    def _write_script_to_file(self, file, script: Script):
        """將單個劇本寫入檔案 - 奶奶對話版本"""
        file.write(f"\n奶奶對話劇本 {script.script_id}: {script.source} → {script.target_slot}\n")
        file.write(f"類型: {script.script_type}\n")
        file.write(f"目標資訊: {', '.join(script.target_info)}\n")
        file.write("-" * 40 + "\n")

        for step in script.steps:
            file.write(f"第{step.step_number}步:\n")
            file.write(f"  👧 孫女: {step.child_dialogue}\n")
            file.write(f"  👵 預期奶奶: {step.expected_grandma_response}\n\n")
        file.write("\n")

def main():
    """主程序 - 奶奶對話專用版本"""
    print("📝 奶奶對話劇本生成器 - Colab版本")
    print("=" * 40)
    print("💡 使用說明:")
    print("1. 請先在代碼中設置您的OpenAI API金鑰")
    print("2. 確保data資料夾中包含以下檔案:")
    print("   data/")
    print("   ├── 01_背景資訊.docx")
    print("   ├── 03_喜好興趣.docx")
    print("   └── 02_作息.xlsx")
    print("3. 本版本專門生成與奶奶對話的劇本")
    print("=" * 40)

    # 檢查Google Drive
    if not os.path.exists("/content/drive/MyDrive"):
        print("❌ Google Drive 未掛載")
        print("💡 請先執行: from google.colab import drive; drive.mount('/content/drive')")
        return

    try:
        # 初始化生成器
        generator = ScriptGenerator()

        # 檢查data資料夾中的必要檔案
        required_files = ["01_背景資訊.docx", "03_喜好興趣.docx", "02_作息.xlsx"]

        for filename in required_files:
            file_path = os.path.join(generator.data_path, filename)
            if not os.path.exists(file_path):
                print(f"❌ 找不到必要檔案: {file_path}")
                print("💡 請確保data資料夾中包含以下檔案:")
                for f in required_files:
                    print(f"   - {f}")
                return

        # 生成所有劇本
        print("\n🎭 開始生成奶奶對話劇本...")

        try:
            scripts = generator.generate_all_scripts()

            # 保存劇本
            script_file = generator.save_scripts(scripts)

            # 統計資訊
            print(f"\n🎉 奶奶對話劇本生成完成！")
            print(f"📊 統計資訊:")
            print(f"  總計劇本: {len(scripts)} 個")
            print(f"  對話對象: 奶奶")
            print(f"\n💾 檔案已保存:")
            print(f"  劇本檔案: {script_file}")
            print(f"\n📁 輸出路徑:")
            print(f"  會話資料夾: {generator.session_folder}")
            print(f"  劇本路徑: {generator.script_path}")

            # 顯示劇本概覽
            sources = set(s.source for s in scripts)
            slots = set(s.target_slot for s in scripts)
            print(f"\n📋 劇本概覽:")
            print(f"  話題來源: {', '.join(sources)}")
            print(f"  目的槽: {', '.join(slots)}")

            # 顯示劇本樣例
            if scripts:
                sample_script = scripts[0]
                print(f"\n📝 奶奶對話劇本樣例 (劇本{sample_script.script_id}):")
                print(f"  來源: {sample_script.source} → 目標: {sample_script.target_slot}")
                print(f"  對話步驟: {sample_script.total_steps}步")
                if sample_script.steps:
                    print(f"  第1步範例:")
                    print(f"    👧 孫女: {sample_script.steps[0].child_dialogue}")
                    print(f"    👵 預期奶奶: {sample_script.steps[0].expected_grandma_response}")

        except Exception as e:
            print(f"❌ 奶奶對話劇本生成失敗: {e}")

            # 檢查是否有部分進度
            if hasattr(generator, 'generated_scripts'):
                total_generated = len(generator.generated_scripts)
                if total_generated > 0:
                    print(f"\n📊 已生成進度:")
                    print(f"  已完成劇本: {total_generated}/12")
                    print(f"\n💡 進度已自動保存，重新執行程式將從中斷處繼續")

            print("\n🔧 可能的解決方案:")
            print("1. 檢查API金鑰是否正確")
            print("2. 檢查網路連線")
            print("3. 檢查OpenAI API額度")
            print("4. 重新執行程式將自動從中斷處繼續")

    except Exception as e:
        print(f"❌ 奶奶對話劇本生成失敗: {e}")
        logger.error(f"奶奶對話劇本生成錯誤: {e}")
        print("\n🔧 可能的解決方案:")
        print("1. 檢查API金鑰是否正確")
        print("2. 檢查網路連線")
        print("3. 檢查OpenAI API額度")
        print("4. 檢查data資料夾中的檔案是否存在")
        print("5. 確認所有劇本都專門與奶奶對話")

if __name__ == "__main__":
    main()
